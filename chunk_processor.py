"""
智能切片处理器
解决复杂内容元素的精准切片问题
支持多级切片策略和语义边界识别
"""

import os
import re
import uuid
from typing import List, Dict, Any, Optional, Tuple
import logging
from datetime import datetime

# 先初始化日志，确保后续导入块可用
logger = logging.getLogger(__name__)

# 尝试导入各种文件解析库
try:
    import pdfplumber
    PDF_PARSER_AVAILABLE = True
except ImportError:
    PDF_PARSER_AVAILABLE = False
    logger.warning("pdfplumber未安装，PDF解析功能受限")

try:
    from docx import Document
    DOCX_PARSER_AVAILABLE = True
except ImportError:
    DOCX_PARSER_AVAILABLE = False
    logger.warning("python-docx未安装，Word解析功能受限")

try:
    import pandas as pd
    EXCEL_PARSER_AVAILABLE = True
except ImportError:
    EXCEL_PARSER_AVAILABLE = False
    logger.warning("pandas未安装，Excel解析功能受限")

try:
    import pytesseract
    from PIL import Image
    OCR_PARSER_AVAILABLE = True
except ImportError:
    OCR_PARSER_AVAILABLE = False
    logger.warning("pytesseract/PIL未安装，图片OCR功能受限")

try:
    from feishu_integration import FeishuDocumentParser, FeishuAPIError
    FEISHU_PARSER_AVAILABLE = True
except ImportError:
    FEISHU_PARSER_AVAILABLE = False
    logger.warning("feishu_integration未安装，飞书文档解析功能受限")

try:
    from xmindparser import xmind_to_dict
    XMIND_PARSER_AVAILABLE = True
except ImportError:
    XMIND_PARSER_AVAILABLE = False
    logger.warning("xmindparser未安装，XMind导图解析功能受限")


class ContentElement:
    """内容元素基类"""
    
    def __init__(self, element_type: str, content: Any, metadata: Dict[str, Any] = None):
        self.id = str(uuid.uuid4())
        self.element_type = element_type  # text, table, image, heading, code
        self.content = content
        self.metadata = metadata or {}
        self.position = metadata.get('position', 0) if metadata else 0
        self.semantic_context = self.metadata.get('semantic_context', '')
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典"""
        return {
            'id': self.id,
            'element_type': self.element_type,
            'content': self.content,
            'metadata': self.metadata,
            'position': self.position,
            'semantic_context': self.semantic_context
        }


class ChunkingStrategy:
    """切片策略基类"""
    
    def __init__(self, min_chunk_size: int = 200, max_chunk_size: int = 1000, overlap_size: int = 50):
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size
        self.overlap_size = overlap_size
    
    def chunk_elements(self, elements: List[ContentElement]) -> List[List[ContentElement]]:
        """将元素分组为切片"""
        raise NotImplementedError("子类必须实现此方法")


class StructureBasedStrategy(ChunkingStrategy):
    """基于文档结构的切片策略"""
    
    def chunk_elements(self, elements: List[ContentElement]) -> List[List[ContentElement]]:
        """基于文档结构分组（标题、章节等）"""
        chunks = []
        current_chunk = []
        current_size = 0
        
        for element in elements:
            element_size = len(str(element.content)) if element.element_type == 'text' else 100
            
            # 如果遇到标题，开始新切片
            if element.element_type == 'heading' and current_chunk:
                if current_size >= self.min_chunk_size:
                    chunks.append(current_chunk)
                    current_chunk = [element]
                    current_size = element_size
                else:
                    current_chunk.append(element)
                    current_size += element_size
            else:
                # 检查是否超过最大大小
                if current_size + element_size > self.max_chunk_size and current_chunk:
                    chunks.append(current_chunk)
                    current_chunk = [element]
                    current_size = element_size
                else:
                    current_chunk.append(element)
                    current_size += element_size
        
        # 添加最后一个切片
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks


class SemanticBasedStrategy(ChunkingStrategy):
    """基于语义边界的切片策略"""
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.sentence_endings = ['.', '。', '!', '！', '?', '？', ';', '；']
    
    def chunk_elements(self, elements: List[ContentElement]) -> List[List[ContentElement]]:
        """基于语义边界分组"""
        chunks = []
        current_chunk = []
        current_text = ""
        
        for element in elements:
            if element.element_type == 'text':
                element_text = str(element.content)
                
                # 查找语义边界
                semantic_boundary = self._find_semantic_boundary(current_text + " " + element_text)
                
                if semantic_boundary and len(current_text) >= self.min_chunk_size:
                    # 当前文本已达到最小大小且有语义边界，创建切片
                    chunks.append(current_chunk.copy())
                    current_chunk = [element]
                    current_text = element_text
                else:
                    # 添加到当前切片
                    current_chunk.append(element)
                    current_text += " " + element_text if current_text else element_text
                    
                    # 检查是否超过最大大小
                    if len(current_text) > self.max_chunk_size:
                        # 强制在最近的句子边界处切割
                        forced_chunks = self._force_chunk_by_sentences(current_chunk, current_text)
                        if len(forced_chunks) > 1:
                            chunks.extend(forced_chunks[:-1])
                            current_chunk = forced_chunks[-1]
                            current_text = self._get_text_from_elements(current_chunk)
                        else:
                            chunks.append(current_chunk)
                            current_chunk = []
                            current_text = ""
            else:
                # 非文本元素，直接添加到当前切片
                current_chunk.append(element)
        
        # 添加最后一个切片
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def _find_semantic_boundary(self, text: str) -> bool:
        """查找语义边界"""
        # 简单的实现：检查是否有句子结束符
        for ending in self.sentence_endings:
            if ending in text[-50:]:  # 检查最后50个字符
                return True
        return False
    
    def _force_chunk_by_sentences(self, elements: List[ContentElement], text: str) -> List[List[ContentElement]]:
        """按句子强制分割"""
        # 简化实现：按近似大小分割
        # 在实际应用中，应该使用NLP模型进行句子分割
        chunk_count = max(1, len(text) // self.max_chunk_size)
        avg_size = len(elements) // chunk_count
        
        chunks = []
        for i in range(0, len(elements), avg_size):
            chunk = elements[i:i + avg_size]
            if chunk:
                chunks.append(chunk)
        
        return chunks
    
    def _get_text_from_elements(self, elements: List[ContentElement]) -> str:
        """从元素中提取文本"""
        text_parts = []
        for element in elements:
            if element.element_type == 'text':
                text_parts.append(str(element.content))
        return " ".join(text_parts)


class HybridChunkingStrategy(ChunkingStrategy):
    """混合切片策略：结合结构和语义"""
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.structure_strategy = StructureBasedStrategy(*args, **kwargs)
        self.semantic_strategy = SemanticBasedStrategy(*args, **kwargs)
    
    def chunk_elements(self, elements: List[ContentElement]) -> List[List[ContentElement]]:
        """混合策略：先按结构，再按语义优化"""
        # 第一步：按结构分组
        structure_chunks = self.structure_strategy.chunk_elements(elements)
        
        # 第二步：对每个结构切片进行语义优化
        final_chunks = []
        for chunk in structure_chunks:
            if self._needs_semantic_split(chunk):
                # 需要进一步语义分割
                semantic_chunks = self.semantic_strategy.chunk_elements(chunk)
                final_chunks.extend(semantic_chunks)
            else:
                final_chunks.append(chunk)
        
        return final_chunks
    
    def _needs_semantic_split(self, chunk: List[ContentElement]) -> bool:
        """检查是否需要语义分割"""
        # 计算文本长度
        text_length = 0
        for element in chunk:
            if element.element_type == 'text':
                text_length += len(str(element.content))
        
        # 如果文本长度超过最大大小的1.5倍，需要分割
        return text_length > self.max_chunk_size * 1.5


class AdvancedChunkingStrategy(ChunkingStrategy):
    """高级切片策略：针对复杂内容元素优化"""
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        # 针对不同元素类型的权重
        self.element_weights = {
            'text': 1,
            'heading': 2,      # 标题更重要
            'table': 3,        # 表格需要保持完整
            'image': 2,        # 图片描述
            'code': 3          # 代码块需要保持完整
        }
        # 复杂度阈值
        self.complexity_threshold = 5  # 每个切片的平均元素复杂度
    
    def chunk_elements(self, elements: List[ContentElement]) -> List[List[ContentElement]]:
        """高级策略：考虑元素类型、语义边界和复杂度"""
        chunks = []
        current_chunk = []
        current_complexity = 0
        current_text_length = 0
        
        for element in elements:
            element_weight = self.element_weights.get(element.element_type, 1)
            element_text_length = len(str(element.content)) if element.element_type == 'text' else 0
            
            # 检查是否需要开始新切片
            new_chunk_needed = False
            
            # 规则1：遇到标题且当前切片已有内容
            if element.element_type == 'heading' and current_chunk:
                new_chunk_needed = True
            
            # 规则2：当前切片复杂度超过阈值
            elif current_complexity + element_weight > self.complexity_threshold * 2:
                new_chunk_needed = True
            
            # 规则3：文本长度超过最大限制
            elif current_text_length + element_text_length > self.max_chunk_size:
                new_chunk_needed = True
            
            # 规则4：表格和代码块应尽量保持独立
            elif element.element_type in ['table', 'code'] and current_chunk:
                # 如果当前切片已经有内容，且不是以表格/代码开始，则新建切片
                if not any(e.element_type in ['table', 'code'] for e in current_chunk):
                    new_chunk_needed = True
            
            if new_chunk_needed and current_chunk:
                chunks.append(current_chunk)
                current_chunk = [element]
                current_complexity = element_weight
                current_text_length = element_text_length
            else:
                current_chunk.append(element)
                current_complexity += element_weight
                current_text_length += element_text_length
        
        # 添加最后一个切片
        if current_chunk:
            chunks.append(current_chunk)
        
        # 后处理：确保没有切片过大
        final_chunks = []
        for chunk in chunks:
            chunk_text_length = sum(len(str(e.content)) for e in chunk if e.element_type == 'text')
            if chunk_text_length > self.max_chunk_size * 1.2:
                # 需要进一步分割
                sub_chunks = self._split_large_chunk(chunk, chunk_text_length)
                final_chunks.extend(sub_chunks)
            else:
                final_chunks.append(chunk)
        
        return final_chunks
    
    def _split_large_chunk(self, chunk: List[ContentElement], text_length: int) -> List[List[ContentElement]]:
        """分割过大的切片"""
        if len(chunk) <= 1:
            return [chunk]
        
        # 计算每个元素的大致文本长度
        element_sizes = []
        for element in chunk:
            if element.element_type == 'text':
                element_sizes.append(len(str(element.content)))
            else:
                element_sizes.append(100)  # 非文本元素的估计大小
        
        # 尝试将切片分为大小相近的两部分
        target_size = text_length // 2
        current_size = 0
        split_index = 0
        
        for i, size in enumerate(element_sizes):
            current_size += size
            if current_size >= target_size:
                split_index = i + 1
                break
        
        if split_index == 0 or split_index >= len(chunk):
            return [chunk]
        
        first_part = chunk[:split_index]
        second_part = chunk[split_index:]
        
        return [first_part, second_part]
    
    def calculate_chunk_quality(self, chunk: List[ContentElement]) -> Dict[str, Any]:
        """计算切片质量评分"""
        text_length = 0
        element_types = {}
        has_heading = False
        has_table = False
        
        for element in chunk:
            element_types[element.element_type] = element_types.get(element.element_type, 0) + 1
            if element.element_type == 'text':
                text_length += len(str(element.content))
            elif element.element_type == 'heading':
                has_heading = True
            elif element.element_type == 'table':
                has_table = True
        
        # 质量评分规则
        score = 0
        if self.min_chunk_size <= text_length <= self.max_chunk_size:
            score += 2
        elif text_length < self.min_chunk_size:
            score -= 1
        else:
            score -= 2
        
        if has_heading:
            score += 1
        if len(element_types) > 1:
            score += 1  # 多样性加分
        if has_table and len(chunk) == 1:
            score += 1  # 表格独立切片
        
        return {
            'score': score,
            'text_length': text_length,
            'element_types': element_types,
            'has_heading': has_heading,
            'has_table': has_table,
            'is_optimal': score >= 3
        }


class FastChunkingStrategy(ChunkingStrategy):
    """快速切片策略：牺牲精度换取速度"""
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        # 固定大小切片，忽略语义边界
        self.fixed_chunk_size = 500  # 目标文本大小
    
    def chunk_elements(self, elements: List[ContentElement]) -> List[List[ContentElement]]:
        """快速策略：按固定文本大小切片"""
        chunks = []
        current_chunk = []
        current_text_length = 0
        
        for element in elements:
            element_text_length = len(str(element.content)) if element.element_type == 'text' else 0
            
            # 如果当前切片已满，开始新切片
            if current_text_length + element_text_length > self.fixed_chunk_size and current_chunk:
                chunks.append(current_chunk)
                current_chunk = [element]
                current_text_length = element_text_length
            else:
                current_chunk.append(element)
                current_text_length += element_text_length
        
        # 添加最后一个切片
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks


class ContentParser:
    """内容解析器基类"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.elements: List[ContentElement] = []
    
    def parse(self) -> List[ContentElement]:
        """解析文件内容"""
        raise NotImplementedError("子类必须实现此方法")
    
    def extract_metadata(self) -> Dict[str, Any]:
        """提取文件元数据"""
        return {
            'file_path': self.file_path,
            'file_size': os.path.getsize(self.file_path) if os.path.exists(self.file_path) else 0,
            'parse_time': datetime.now().isoformat()
        }


class PDFParser(ContentParser):
    """PDF文件解析器"""

    def parse(self) -> List[ContentElement]:
        logger.info(f"开始解析PDF文件: {self.file_path}, 方法: pdfplumber")
        if not PDF_PARSER_AVAILABLE:
            logger.warning("PDF解析器不可用，返回错误提示切片")
            return [ContentElement('text', '无法解析文件：pdfplumber未安装，PDF解析功能不可用', {'error': 'PDF_PARSER_AVAILABLE=False', 'position': 0})]

        elements = []

        try:
            with pdfplumber.open(self.file_path) as pdf:
                logger.info(f"  PDF共{len(pdf.pages)}页")
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        paragraphs = re.split(r'\n\s*\n', text)
                        for para_num, paragraph in enumerate(paragraphs):
                            if paragraph.strip():
                                elements.append(ContentElement('text', paragraph.strip(), {
                                    'page': page_num + 1, 'paragraph': para_num + 1, 'position': len(elements)
                                }))

                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables):
                        if table:
                            md_lines = []
                            for row_idx, row in enumerate(table):
                                cells = [str(c).strip() if c else '' for c in row]
                                md_lines.append('| ' + ' | '.join(cells) + ' |')
                                if row_idx == 0:
                                    md_lines.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
                            elements.append(ContentElement('table', '\n'.join(md_lines), {
                                'page': page_num + 1, 'table': table_num + 1,
                                'rows': len(table), 'columns': len(table[0]) if table else 0,
                                'position': len(elements)
                            }))

                    # 检测页面中的图片/流程图（pdfplumber page.images）
                    if hasattr(page, 'images') and page.images:
                        for img_idx, img in enumerate(page.images):
                            # 获取图片周围的文本作为上下文
                            img_x0, img_y0 = img.get('x0', 0), img.get('y0', 0)
                            img_x1, img_y1 = img.get('x1', 0), img.get('y1', 0)
                            nearby_text = ''
                            if hasattr(page, 'chars') and page.chars:
                                nearby = [c['text'] for c in page.chars
                                          if abs(c.get('top', 0) - img_y1) < 50
                                          and c.get('text', '').strip()]
                                if nearby:
                                    nearby_text = ' '.join(nearby[-10:])[:200]
                            alt_text = nearby_text if nearby_text else '[流程图/图片]'
                            elements.append(ContentElement('image', alt_text, {
                                'page': page_num + 1, 'image': img_idx + 1,
                                'bbox': [img_x0, img_y0, img_x1, img_y1],
                                'width': img.get('width', 0), 'height': img.get('height', 0),
                                'position': len(elements)
                            }))

            if not elements:
                logger.warning("PDF文件无有效文本内容，返回提示切片")
                return [ContentElement('text', '无法解析文件：PDF无有效文本内容（可能为扫描件/图片型PDF）', {'error': 'no_text_content', 'position': 0})]

            logger.info(f"PDF解析完成: {len(elements)} 个元素, {len(pdf.pages)}页")
            return elements

        except Exception as e:
            logger.error(f"PDF解析失败: {e}", exc_info=True)
            return [ContentElement('text', f'无法解析文件：PDF解析出错 ({e})', {'error': str(e), 'position': 0})]


class DocxParser(ContentParser):
    """Word文档解析器 - 支持段落/表格/文本框/页眉页脚/图片降级"""

    def parse(self) -> List[ContentElement]:
        logger.info(f"开始解析Word文档: {self.file_path}, 方法: python-docx")
        if not DOCX_PARSER_AVAILABLE:
            logger.warning("Word解析器不可用，返回错误提示切片")
            return [ContentElement('text', '无法解析文件：python-docx未安装，Word解析功能不可用', {'error': 'DOCX_PARSER_AVAILABLE=False', 'position': 0})]

        elements = []
        stat = {'paragraphs': 0, 'tables': 0, 'headers': 0, 'footers': 0, 'textboxes': 0, 'images': 0}

        try:
            doc = Document(self.file_path)

            # --- 1. 提取页眉 ---
            for sec_idx, section in enumerate(doc.sections):
                try:
                    header = section.header
                    if header and header.paragraphs:
                        header_texts = [p.text.strip() for p in header.paragraphs if p.text.strip()]
                        if header_texts:
                            text = ' | '.join(header_texts)
                            elements.append(ContentElement('text', text, {
                                'source': 'header', 'section': sec_idx + 1, 'position': len(elements)
                            }))
                            stat['headers'] += 1
                except Exception:
                    pass

            # --- 2. 提取段落（含文本框内容） ---
            for para_num, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    style_name = paragraph.style.name if paragraph.style else 'Normal'
                    is_heading = style_name.startswith('Heading')
                    element_type = 'heading' if is_heading else 'text'

                    elements.append(ContentElement(element_type, text, {
                        'paragraph': para_num + 1, 'style': style_name, 'is_heading': is_heading,
                        'heading_level': int(style_name.split()[-1]) if is_heading and style_name.split()[-1].isdigit() else 0,
                        'source': 'paragraph', 'position': len(elements)
                    }))
                    stat['paragraphs'] += 1

            # --- 3. 提取文本框（通过XML遍历w:txbxContent） ---
            try:
                from lxml import etree
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                for txbx in doc.element.body.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent'):
                    texts = []
                    for t in txbx.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                        if t.text:
                            texts.append(t.text)
                    full = ''.join(texts).strip()
                    if full:
                        elements.append(ContentElement('text', full, {
                            'source': 'textbox', 'position': len(elements)
                        }))
                        stat['textboxes'] += 1
            except Exception:
                pass

            # --- 4. 提取表格（转Markdown格式） ---
            for table_num, table in enumerate(doc.tables):
                md_lines = []
                for row_idx, row in enumerate(table.rows):
                    cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    md_lines.append('| ' + ' | '.join(cells) + ' |')
                    if row_idx == 0:
                        md_lines.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
                if md_lines:
                    elements.append(ContentElement('table', '\n'.join(md_lines), {
                        'table': table_num + 1, 'rows': len(table.rows),
                        'columns': len(table.rows[0].cells) if table.rows else 0,
                        'source': 'table', 'position': len(elements)
                    }))
                    stat['tables'] += 1

            # --- 5. 提取图片（alt文本 → 周围100字符 → [图片]） ---
            try:
                for para_idx, para in enumerate(doc.paragraphs):
                    for run in para.runs:
                        if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                            alt_text = ''
                            # 优先取 docPr 的 name/descr 作为替代文本
                            for docPr in run._element.iter('{http://schemas.openxmlformats.org/drawingml/2006/main}docPr'):
                                alt_text = docPr.get('descr') or docPr.get('name') or ''
                            if not alt_text:
                                # 取图片周围100字符
                                context_start = max(0, para_idx - 3)
                                context_end = min(len(doc.paragraphs), para_idx + 4)
                                context_parts = []
                                for cp in range(context_start, context_end):
                                    t = doc.paragraphs[cp].text.strip()
                                    if t:
                                        context_parts.append(t)
                                context_str = ' '.join(context_parts)
                                alt_text = context_str[:100] if context_str else '[图片]'
                            elements.append(ContentElement('image', alt_text, {
                                'source': 'image', 'paragraph': para_idx + 1, 'position': len(elements)
                            }))
                            stat['images'] += 1
            except Exception:
                pass

            # --- 6. 提取页脚 ---
            for sec_idx, section in enumerate(doc.sections):
                try:
                    footer = section.footer
                    if footer and footer.paragraphs:
                        footer_texts = [p.text.strip() for p in footer.paragraphs if p.text.strip()]
                        if footer_texts:
                            text = ' | '.join(footer_texts)
                            elements.append(ContentElement('text', text, {
                                'source': 'footer', 'section': sec_idx + 1, 'position': len(elements)
                            }))
                            stat['footers'] += 1
                except Exception:
                    pass

            # --- 7. 兜底检查: 提取文本总长度 < 50 则返回错误切片 ---
            total_text_len = sum(len(str(e.content)) for e in elements if e.element_type in ('text', 'heading', 'table'))
            if total_text_len < 50:
                logger.warning(f"Word文档有效内容过短({total_text_len}字符)，返回错误提示切片")
                return [ContentElement('text', '无法解析文件：Word文档内容为空或仅有少量非文字内容', {'error': 'content_too_short', 'total_text_length': total_text_len, 'position': 0})]

            logger.info(f"Word文档解析完成: {len(elements)} 个元素 (段落{stat['paragraphs']}/表格{stat['tables']}/页眉{stat['headers']}/页脚{stat['footers']}/文本框{stat['textboxes']}/图片{stat['images']})")
            return elements

        except Exception as e:
            logger.error(f"Word文档解析失败: {e}", exc_info=True)
            return [ContentElement('text', f'无法解析文件：Word文档解析出错 ({e})', {'error': str(e), 'position': 0})]


class ExcelParser(ContentParser):
    """Excel文件解析器 - 自适应多Sheet全量逐行切片"""

    def parse(self) -> List[ContentElement]:
        file_ext = os.path.splitext(self.file_path)[1].lower()
        logger.info(f"开始解析Excel文件: {self.file_path}, 方法: {'csv模块' if file_ext == '.csv' else 'openpyxl'}")
        try:
            import openpyxl
        except ImportError:
            logger.warning("openpyxl未安装")
            return [ContentElement('text', '无法解析文件：openpyxl未安装，Excel/CSV解析功能不可用', {'error': 'openpyxl not installed', 'position': 0})]

        if file_ext == '.csv':
            try:
                elements = self._parse_csv()
                if elements:
                    logger.info(f"CSV文件解析完成: {len(elements)} 个元素")
                    sheet_list = 'CSV文件共1个工作表: 数据'
                    elements.insert(0, ContentElement('text', sheet_list, {
                        'is_sheet_list': True, 'sheet_names': ['数据'], 'position': 0
                    }))
                    return elements
            except Exception as e:
                logger.warning(f"CSV解析失败: {e}, 降级尝试openpyxl")

        elements = []
        KEYWORDS = ['项目', '阶段', '进度', '负责人', 'DPM', '部门', '计划名称']
        sheet_names = []

        try:
            wb = openpyxl.load_workbook(self.file_path, read_only=False)
            sheet_names = wb.sheetnames
        except Exception as e:
            logger.error(f"无法打开Excel文件: {e}", exc_info=True)
            return [ContentElement('text', f'无法解析文件：Excel文件损坏或格式不支持 ({e})', {'error': str(e), 'position': 0})]

        for sheet_name in sheet_names:
            logger.info(f"  处理工作表: '{sheet_name}'")
            try:
                ws = wb[sheet_name]
            except Exception:
                elements.append(ContentElement('text', f'工作表【{sheet_name}】读取失败', {
                    'sheet_name': sheet_name, 'error': 'sheet_read_error', 'position': len(elements)
                }))
                continue

            # --- 收集所有行数据（分批读取以防大文件OOM） ---
            all_rows = []
            for row_idx, row in enumerate(ws.iter_rows(values_only=False), start=1):
                row_values = []
                is_empty = True
                for cell in row:
                    v = cell.value
                    if v is not None:
                        is_empty = False
                        row_values.append(str(v).strip())
                    else:
                        row_values.append('')
                all_rows.append((row_idx, row_values, is_empty))

            # --- 空sheet检测 ---
            if not all_rows or all(r[2] for r in all_rows):
                elements.append(ContentElement('text', f'工作表【{sheet_name}】无有效数据', {
                    'sheet_name': sheet_name, 'empty_sheet': True, 'position': len(elements)
                }))
                continue

            # --- 自动定位数据起始行（扫描前20行找关键词行） ---
            header_row_idx = None
            header_columns = []
            scan_limit = min(20, len(all_rows))

            for ri in range(scan_limit):
                row_num, row_vals, is_empty = all_rows[ri]
                row_text = ' '.join(v for v in row_vals if v)
                if any(kw in row_text for kw in KEYWORDS):
                    header_row_idx = ri
                    header_columns = [v for v in row_vals]
                    # 检测合并表头：如果当前行有空单元格且下一行非空，降级到下一行
                    if ri + 1 < scan_limit:
                        next_row_num, next_row_vals, next_is_empty = all_rows[ri + 1]
                        empty_count = sum(1 for v in header_columns if not v.strip())
                        if empty_count > 0 and not next_is_empty:
                            header_row_idx = ri + 1
                            header_columns = [v for v in next_row_vals]
                            logger.info(f"    检测到合并表头（{empty_count}个空单元格），降级到行{next_row_num}: {header_columns}")
                    logger.info(f"    在行{row_num}检测到关键词表头: {header_columns}")
                    break

            if header_row_idx is None:
                for ri in range(scan_limit):
                    row_num, row_vals, is_empty = all_rows[ri]
                    if not is_empty:
                        header_row_idx = ri
                        header_columns = [v for v in row_vals]
                        logger.info(f"    未检测到关键词，使用第一个非空行{row_num}作为列名: {header_columns}")
                        break

            if header_row_idx is None:
                header_row_idx = 0
                header_columns = [f'列{i+1}' for i in range(len(all_rows[0][1]))]
                logger.info("    所有前20行均为空，使用默认列名")

            # --- 列名去重 ---
            deduped_headers = []
            seen = {}
            for h in header_columns:
                h_clean = h if h else f'列{len(deduped_headers)+1}'
                if h_clean in seen:
                    seen[h_clean] += 1
                    deduped_headers.append(f'{h_clean}_{seen[h_clean]}')
                else:
                    seen[h_clean] = 0
                    deduped_headers.append(h_clean)

            # --- 描述切片：字段列表 + 有效行数 ---
            data_rows = [r for r in all_rows[header_row_idx + 1:] if not r[2]]
            desc = f'工作表【{sheet_name}】字段: {", ".join(deduped_headers)} | 有效数据行数: {len(data_rows)}'
            elements.append(ContentElement('text', desc, {
                'sheet_name': sheet_name, 'is_description': True,
                'fields': deduped_headers, 'data_row_count': len(data_rows), 'position': len(elements)
            }))

            # --- 全量逐行切片 ---
            for ri in range(header_row_idx + 1, len(all_rows)):
                row_num, row_vals, is_empty = all_rows[ri]
                if is_empty:
                    continue
                field_parts = []
                for ci in range(min(len(deduped_headers), len(row_vals))):
                    val = row_vals[ci]
                    if val:
                        field_parts.append(f'{deduped_headers[ci]}: {val}')
                if field_parts:
                    content = f'【工作表：{sheet_name}】第{row_num}行\n' + ' | '.join(field_parts)
                    elements.append(ContentElement('text', content, {
                        'sheet_name': sheet_name, 'excel_row': row_num,
                        'row_index_in_sheet': ri - header_row_idx,
                        'position': len(elements)
                    }))

            logger.info(f"    工作表'{sheet_name}'生成 {len(data_rows)} 个数据切片")

        wb.close()

        # --- 整个文件的工作表列表切片 ---
        sheet_list = f'Excel文件共{len(sheet_names)}个工作表: {", ".join(sheet_names)}'
        elements.insert(0, ContentElement('text', sheet_list, {
            'is_sheet_list': True, 'sheet_names': sheet_names, 'position': 0
        }))

        logger.info(f"Excel文件解析完成: {len(elements)} 个元素, {len(sheet_names)}个工作表")
        return elements

    def _parse_csv(self) -> List[ContentElement]:
        """使用csv模块解析CSV文件"""
        import csv
        elements = []
        try:
            with open(self.file_path, 'r', encoding='utf-8-sig') as f:
                reader = csv.reader(f)
                all_rows = []
                for row_idx, row in enumerate(reader, start=1):
                    row_values = [v.strip() for v in row]
                    is_empty = all(not v for v in row_values)
                    all_rows.append((row_idx, row_values, is_empty))
        except UnicodeDecodeError:
            with open(self.file_path, 'r', encoding='gbk') as f:
                reader = csv.reader(f)
                all_rows = []
                for row_idx, row in enumerate(reader, start=1):
                    row_values = [v.strip() for v in row]
                    is_empty = all(not v for v in row_values)
                    all_rows.append((row_idx, row_values, is_empty))
        except Exception as e:
            logger.error(f"CSV文件读取失败: {e}", exc_info=True)
            return [ContentElement('text', f'无法解析文件：CSV读取失败 ({e})', {'error': str(e), 'position': 0})]

        if not all_rows or all(r[2] for r in all_rows):
            return [ContentElement('text', 'CSV文件无有效数据', {'empty_sheet': True, 'position': 0})]

        KEYWORDS = ['项目', '阶段', '进度', '负责人', 'DPM', '部门', '计划名称']
        scan_limit = min(20, len(all_rows))
        header_row_idx = None
        header_columns = []

        for ri in range(scan_limit):
            row_num, row_vals, is_empty = all_rows[ri]
            row_text = ' '.join(v for v in row_vals if v)
            if any(kw in row_text for kw in KEYWORDS):
                header_row_idx = ri
                header_columns = [v for v in row_vals]
                if ri + 1 < scan_limit:
                    next_row_num, next_row_vals, next_is_empty = all_rows[ri + 1]
                    empty_count = sum(1 for v in header_columns if not v.strip())
                    if empty_count > 0 and not next_is_empty:
                        header_row_idx = ri + 1
                        header_columns = [v for v in next_row_vals]
                        logger.info(f"    CSV检测到合并表头，降级到行{next_row_num}: {header_columns}")
                logger.info(f"    在行{row_num}检测到关键词表头: {header_columns}")
                break

        if header_row_idx is None:
            for ri in range(scan_limit):
                row_num, row_vals, is_empty = all_rows[ri]
                if not is_empty:
                    header_row_idx = ri
                    header_columns = [v for v in row_vals]
                    logger.info(f"    未检测到关键词，使用第一个非空行{row_num}作为列名: {header_columns}")
                    break

        if header_row_idx is None:
            header_row_idx = 0
            header_columns = [f'列{i+1}' for i in range(len(all_rows[0][1]))]

        deduped_headers = []
        seen = {}
        for h in header_columns:
            h_clean = h if h else f'列{len(deduped_headers)+1}'
            if h_clean in seen:
                seen[h_clean] += 1
                deduped_headers.append(f'{h_clean}_{seen[h_clean]}')
            else:
                seen[h_clean] = 0
                deduped_headers.append(h_clean)

        data_rows = [r for r in all_rows[header_row_idx + 1:] if not r[2]]
        desc = f'CSV文件字段: {", ".join(deduped_headers)} | 有效数据行数: {len(data_rows)}'
        elements.append(ContentElement('text', desc, {
            'sheet_name': '数据', 'is_description': True,
            'fields': deduped_headers, 'data_row_count': len(data_rows), 'position': len(elements)
        }))

        for ri in range(header_row_idx + 1, len(all_rows)):
            row_num, row_vals, is_empty = all_rows[ri]
            if is_empty:
                continue
            field_parts = []
            for ci in range(min(len(deduped_headers), len(row_vals))):
                val = row_vals[ci]
                if val:
                    field_parts.append(f'{deduped_headers[ci]}: {val}')
            if field_parts:
                content = f'第{row_num}行\n' + ' | '.join(field_parts)
                elements.append(ContentElement('text', content, {
                    'excel_row': row_num,
                    'row_index_in_sheet': ri - header_row_idx,
                    'position': len(elements)
                }))

        return elements

    def extract_metadata(self) -> Dict[str, Any]:
        meta = super().extract_metadata()
        try:
            import openpyxl
            wb = openpyxl.load_workbook(self.file_path, read_only=True)
            meta['sheet_count'] = len(wb.sheetnames)
            meta['sheet_names'] = wb.sheetnames
            wb.close()
        except Exception:
            pass
        return meta


class TextParser(ContentParser):
    """文本文件解析器 - 支持Markdown代码块检测"""

    def parse(self) -> List[ContentElement]:
        logger.info(f"开始解析文本文件: {self.file_path}, 方法: UTF-8文本")
        elements = []

        try:
            encodings = ['utf-8', 'utf-8-sig', 'gbk', 'gb18030', 'latin-1']
            content = None
            for enc in encodings:
                try:
                    with open(self.file_path, 'r', encoding=enc) as f:
                        content = f.read()
                    logger.info(f"    使用编码{enc}成功读取")
                    break
                except UnicodeDecodeError:
                    continue
            if content is None:
                raise UnicodeDecodeError(f"无法用任何编码打开文件: {encodings}")

            # 检测是否为Markdown文件
            is_markdown = self.file_path.lower().endswith('.md')

            if is_markdown:
                # Markdown解析：按代码块 → 表格 → 标题 → 段落顺序
                elements = self._parse_markdown(content)
            else:
                # 纯文本解析：按段落分割
                paragraphs = re.split(r'\n\s*\n', content)
                for para_num, paragraph in enumerate(paragraphs):
                    if paragraph.strip():
                        elements.append(ContentElement('text', paragraph.strip(), {
                            'paragraph': para_num + 1, 'position': len(elements)
                        }))

            if not elements:
                logger.warning("文本文件内容为空")
                return [ContentElement('text', '无法解析文件：文本文件内容为空', {'error': 'empty_file', 'position': 0})]

            logger.info(f"文本文件解析完成: {len(elements)} 个元素 ({'Markdown' if is_markdown else '纯文本'})")
            return elements

        except Exception as e:
            logger.error(f"文本文件解析失败: {e}", exc_info=True)
            return [ContentElement('text', f'无法解析文件：文本文件读取失败 ({e})', {'error': str(e), 'position': 0})]

    def _parse_markdown(self, content: str) -> List[ContentElement]:
        """解析Markdown内容，识别代码块、表格、标题等"""
        elements = []

        # 使用正则按块分割：先提取代码块，再处理剩余内容
        remaining = content

        while remaining:
            # 查找下一个代码块或表格
            code_match = re.search(r'```(\w*)\n(.*?)```', remaining, re.DOTALL)
            table_match = re.search(r'^\|(.+)\|(?:\n\|[-| ]+\|)?((?:\n\|.+\|)*)', remaining, re.MULTILINE)
            heading_match = re.search(r'^(#{1,6})\s+(.+)$', remaining, re.MULTILINE)

            # 找到最先出现的特殊块
            matches = []
            if code_match:
                matches.append(('code', code_match.start(), code_match))
            if table_match:
                matches.append(('table', table_match.start(), table_match))
            if heading_match:
                matches.append(('heading', heading_match.start(), heading_match))

            if not matches:
                # 没有特殊块，剩余内容全部作为文本处理
                text_blocks = re.split(r'\n\s*\n', remaining.strip())
                for para_num, para in enumerate(text_blocks):
                    if para.strip():
                        # 检查是否为分隔线（--- 或 ***）
                        if re.match(r'^[-*_]{3,}\s*$', para.strip()):
                            continue
                        elements.append(ContentElement('text', para.strip(), {
                            'position': len(elements)
                        }))
                break

            # 排序，处理最早出现的匹配
            matches.sort(key=lambda x: x[1])
            first_type, first_pos, first_match = matches[0]

            # 处理匹配之前的文本内容
            prefix = remaining[:first_pos].strip()
            if prefix:
                text_blocks = re.split(r'\n\s*\n', prefix)
                for para in text_blocks:
                    if para.strip():
                        if re.match(r'^[-*_]{3,}\s*$', para.strip()):
                            continue
                        # 检查是否包含内联标题
                        inline_heading = re.match(r'^(#{1,6})\s+(.+)$', para.strip(), re.MULTILINE)
                        if inline_heading:
                            elements.append(ContentElement('heading', inline_heading.group(2).strip(), {
                                'heading_level': len(inline_heading.group(1)),
                                'position': len(elements)
                            }))
                        else:
                            elements.append(ContentElement('text', para.strip(), {
                                'position': len(elements)
                            }))

            if first_type == 'code':
                lang = first_match.group(1) or 'text'
                code_content = first_match.group(2)
                elements.append(ContentElement('code', code_content, {
                    'language': lang,
                    'position': len(elements)
                }))
                remaining = remaining[first_match.end():]

            elif first_type == 'table':
                table_text = first_match.group(0)
                elements.append(ContentElement('table', table_text.strip(), {
                    'position': len(elements)
                }))
                remaining = remaining[first_match.end():]

            elif first_type == 'heading':
                level = len(first_match.group(1))
                heading_text = first_match.group(2).strip()
                elements.append(ContentElement('heading', heading_text, {
                    'heading_level': level,
                    'position': len(elements)
                }))
                remaining = remaining[first_match.end():]

        return elements


class ImageParser(ContentParser):
    """图片文件解析器（OCR）"""

    def parse(self) -> List[ContentElement]:
        logger.info(f"开始解析图片文件: {self.file_path}, 方法: OCR")
        if not OCR_PARSER_AVAILABLE:
            logger.warning("OCR解析器不可用，返回错误提示切片")
            return [ContentElement('text', '无法解析文件：pytesseract/PIL未安装，图片OCR功能不可用', {'error': 'OCR_PARSER_AVAILABLE=False', 'position': 0})]

        elements = []

        try:
            image = Image.open(self.file_path)
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')

            if text.strip():
                paragraphs = re.split(r'\n\s*\n', text)
                for para_num, paragraph in enumerate(paragraphs):
                    if paragraph.strip():
                        elements.append(ContentElement('text', paragraph.strip(), {
                            'source': 'ocr', 'paragraph': para_num + 1, 'position': len(elements), 'image_size': image.size
                        }))

            if not elements:
                logger.warning("图片OCR未识别到文字")
                return [ContentElement('text', '无法解析文件：图片OCR未识别到文字（可能为空白图片或OCR模型不支持）', {'error': 'ocr_no_text', 'position': 0})]

            logger.info(f"图片文件解析完成: {len(elements)} 个元素")
            return elements

        except Exception as e:
            logger.error(f"图片文件解析失败: {e}", exc_info=True)
            return [ContentElement('text', f'无法解析文件：图片解析出错 ({e})', {'error': str(e), 'position': 0})]


class XMindParser(ContentParser):
    """XMind思维导图文件解析器"""

    def parse(self) -> List[ContentElement]:
        logger.info(f"开始解析XMind导图: {self.file_path}, 方法: xmindparser")
        if not XMIND_PARSER_AVAILABLE:
            logger.warning("XMind解析器不可用，返回错误提示切片")
            return [ContentElement('text', '无法解析文件：xmindparser未安装，XMind导图解析功能不可用', {'error': 'XMIND_PARSER_AVAILABLE=False', 'position': 0})]

        elements = []

        try:
            raw = xmind_to_dict(self.file_path)
            for sheet_idx, topic_data in enumerate(raw):
                sheet_name = topic_data.get('title', f'第{sheet_idx+1}页')
                self._walk_topic(topic_data, elements, sheet_name, 0)

            if not elements:
                logger.warning("XMind导图无有效节点")
                return [ContentElement('text', '无法解析文件：XMind导图无有效节点内容', {'error': 'no_topics', 'position': 0})]

            logger.info(f"XMind导图解析完成: {len(elements)} 个元素")
            return elements

        except Exception as e:
            logger.error(f"XMind导图解析失败: {e}", exc_info=True)
            return [ContentElement('text', f'无法解析文件：XMind导图解析出错 ({e})', {'error': str(e), 'position': 0})]

    def _walk_topic(self, node: dict, elements: list, sheet: str, depth: int):
        """递归遍历导图主题树"""
        title = node.get('title', '').strip()
        if not title:
            return

        element_type = 'heading' if depth <= 1 else 'text'
        meta = {
            'sheet_name': sheet,
            'depth': depth,
            'position': len(elements),
            'source_type': 'xmind_mindmap',
            'chunk_reason': f'思维导图节点(depth={depth})，来自"{sheet}"'
        }

        if depth == 0:
            meta['is_root'] = True
            meta['chunk_reason'] = f'思维导图根节点: {title}'

        elements.append(ContentElement(element_type, title, meta))

        for child in node.get('topics', []):
            self._walk_topic(child, elements, sheet, depth + 1)

    def extract_metadata(self) -> Dict[str, Any]:
        meta = super().extract_metadata()
        meta.update({'source_type': 'xmind_mindmap'})
        return meta


class FeishuParser(ContentParser):
    """飞书云文档解析器"""
    
    def __init__(self, file_path: str):
        """
        初始化飞书解析器
        
        Args:
            file_path: 飞书文档URL或文档ID
        """
        super().__init__(file_path)
        self.parser = None
        self.doc_id = None
        
        if FEISHU_PARSER_AVAILABLE:
            try:
                from feishu_integration import FeishuDocumentParser
                self.parser = FeishuDocumentParser()
                # 从URL提取文档ID
                self.doc_id = self.parser.extract_doc_id_from_url(file_path)
            except ImportError:
                logger.warning("feishu_integration导入失败")
    
    def parse(self) -> List[ContentElement]:
        logger.info(f"开始解析飞书文档: {self.file_path}, 方法: Feishu API")
        if not FEISHU_PARSER_AVAILABLE or not self.parser:
            logger.warning("飞书解析器不可用")
            return [ContentElement('text', '无法解析文件：飞书解析功能不可用', {'error': 'FEISHU_PARSER_AVAILABLE=False', 'position': 0})]

        try:
            permissions = self.parser.check_permission(self.doc_id)
            if not permissions.get('can_read', False):
                logger.error(f"没有权限读取飞书文档: {self.doc_id}")
                return [ContentElement('text', f'无法解析文件：没有权限读取飞书文档', {'error': 'no_permission', 'doc_id': self.doc_id, 'position': 0})]

            elements_data = self.parser.parse_document_to_elements(self.doc_id)
            elements = []
            for elem_data in elements_data:
                elements.append(ContentElement(
                    element_type=elem_data['element_type'],
                    content=elem_data['content'],
                    metadata=elem_data['metadata']
                ))

            if not elements:
                logger.warning("飞书文档解析结果为空")
                return [ContentElement('text', '无法解析文件：飞书文档内容为空', {'error': 'empty_content', 'position': 0})]

            logger.info(f"飞书文档解析完成: {self.doc_id}, {len(elements)} 个元素")
            return elements

        except Exception as e:
            logger.error(f"飞书文档解析失败: {e}", exc_info=True)
            return [ContentElement('text', f'无法解析文件：飞书文档解析出错 ({e})', {'error': str(e), 'position': 0})]
    
    def extract_metadata(self) -> Dict[str, Any]:
        """提取飞书文档元数据"""
        metadata = super().extract_metadata()
        
        if self.doc_id and FEISHU_PARSER_AVAILABLE and self.parser:
            try:
                # 获取文档信息
                doc_info = self.parser.get_document_info(self.file_path)
                metadata.update({
                    'doc_id': self.doc_id,
                    'feishu_info': doc_info,
                    'source': 'feishu',
                    'parse_method': 'api'
                })
            except Exception as e:
                logger.warning(f"获取飞书文档信息失败: {e}")
                metadata.update({
                    'doc_id': self.doc_id,
                    'source': 'feishu',
                    'parse_method': 'api',
                    'error': str(e)
                })
        else:
            metadata.update({
                'doc_id': self.doc_id,
                'source': 'feishu',
                'parse_method': 'unavailable',
                'warning': '飞书解析器不可用'
            })
        
        return metadata


class ChunkProcessor:
    """智能切片处理器"""
    
    def __init__(self, strategy: str = 'hybrid'):
        """
        初始化切片处理器
        
        Args:
            strategy: 切片策略 ('structure', 'semantic', 'hybrid', 'advanced', 'fast')
        """
        self.strategies = {
            'structure': StructureBasedStrategy(),
            'semantic': SemanticBasedStrategy(),
            'hybrid': HybridChunkingStrategy(),
            'advanced': AdvancedChunkingStrategy(),
            'fast': FastChunkingStrategy()
        }
        
        self.strategy = self.strategies.get(strategy, self.strategies['hybrid'])
        
        # 文件解析器映射
        self.parsers = {
            'pdf': PDFParser if PDF_PARSER_AVAILABLE else None,
            'docx': DocxParser if DOCX_PARSER_AVAILABLE else None,
            'xlsx': ExcelParser,
            'csv': ExcelParser,  # CSV也可以用pandas解析
            'xmind': XMindParser if XMIND_PARSER_AVAILABLE else None,
            'txt': TextParser,
            'md': TextParser,  # Markdown暂时按文本处理
            'jpg': ImageParser if OCR_PARSER_AVAILABLE else None,
            'png': ImageParser if OCR_PARSER_AVAILABLE else None,
            'feishu': FeishuParser if FEISHU_PARSER_AVAILABLE else None
        }
    
    def process_file(self, file_path: str, file_type: str) -> Tuple[List[List[ContentElement]], Dict[str, Any]]:
        """
        处理文件：解析 + 切片
        
        Returns:
            tuple: (切片列表, 元数据)
        """
        # 获取解析器
        parser_class = self.parsers.get(file_type.lower())
        if not parser_class:
            logger.warning(f"不支持的文件类型: {file_type}")
            elements = [ContentElement('text', f'无法解析文件：不支持的文件类型({file_type})', {'error': f'unsupported_type:{file_type}', 'position': 0})]
            metadata = {'file_path': file_path, 'fallback': True}
            return [[elements[0]]], metadata

        # 解析文件
        parser = parser_class(file_path)
        elements = parser.parse()
        metadata = parser.extract_metadata()
        
        if not elements:
            logger.warning(f"文件解析结果为空: {file_path}")
            elements = [ContentElement('text', f'无法解析文件：未提取到有效内容', {'error': 'no_elements', 'position': 0})]
            metadata['fallback'] = True
        
        # 执行切片
        chunks = self.strategy.chunk_elements(elements)
        
        # 添加切片统计信息到元数据
        metadata.update({
            'element_count': len(elements),
            'chunk_count': len(chunks),
            'strategy': self.strategy.__class__.__name__,
            'processing_time': datetime.now().isoformat()
        })
        
        logger.info(f"文件处理完成: {len(elements)} 个元素 -> {len(chunks)} 个切片")
        
        return chunks, metadata
    
    def chunk_to_text(self, chunk: List[ContentElement]) -> str:
        """将切片转换为文本（包含切片依据说明）"""
        text_parts = []
        chunk_reasons = []
        
        for element in chunk:
            # 收集切片依据说明
            if 'chunk_reason' in element.metadata:
                reason = element.metadata['chunk_reason']
                if reason and reason not in chunk_reasons:
                    chunk_reasons.append(reason)
            
            if element.element_type == 'text':
                text_parts.append(str(element.content))
            elif element.element_type == 'table':
                # 将表格转换为文本表示
                table_text = self._table_to_text(element.content)
                # 添加表格元数据信息
                table_meta = []
                if 'sheet_name' in element.metadata:
                    table_meta.append(f"工作表: {element.metadata['sheet_name']}")
                if 'rows' in element.metadata and 'columns' in element.metadata:
                    table_meta.append(f"大小: {element.metadata['rows']}行×{element.metadata['columns']}列")
                if 'header_row' in element.metadata and element.metadata['header_row'] is not None:
                    table_meta.append(f"表头行: {element.metadata['header_row']+1}")
                
                meta_str = f" ({', '.join(table_meta)})" if table_meta else ""
                text_parts.append(f"[表格{meta_str}]: {table_text}")
            elif element.element_type == 'heading':
                text_parts.append(f"# {element.content}")
            elif element.element_type == 'image':
                text_parts.append(f"[图片]: {element.content}")
            elif element.element_type == 'code':
                lang = element.metadata.get('language', '') if element.metadata else ''
                code_text = str(element.content)
                text_parts.append(f"[代码块{ ' (' + lang + ')' if lang else '' }]:\n{code_text}")
        
        # 在文本开头添加切片依据说明
        if chunk_reasons:
            reason_text = "切片依据: " + "; ".join(chunk_reasons)
            text_parts.insert(0, reason_text)
        else:
            # 默认切片依据说明
            element_types = set(e.element_type for e in chunk)
            element_count = len(chunk)
            default_reason = f"基于 {element_count} 个元素（类型: {', '.join(element_types)}）自动切片"
            text_parts.insert(0, f"切片依据: {default_reason}")
        
        return "\n\n".join(text_parts)
    
    def _table_to_text(self, table_data) -> str:
        """将表格数据转换为文本（支持Markdown字符串或List[List[str]]）"""
        if not table_data:
            return "空表格"
        if isinstance(table_data, str):
            return table_data
        table_text = []
        for row in table_data:
            row_text = " | ".join(str(cell) for cell in row)
            table_text.append(row_text)
        return "\n".join(table_text)
    
    def analyze_chunk_quality(self, chunk: List[ContentElement]) -> Dict[str, Any]:
        """分析切片质量"""
        text_length = 0
        element_types = {}
        
        for element in chunk:
            element_types[element.element_type] = element_types.get(element.element_type, 0) + 1
            if element.element_type == 'text':
                text_length += len(str(element.content))
        
        return {
            'element_count': len(chunk),
            'text_length': text_length,
            'element_types': element_types,
            'is_optimal': self.strategy.min_chunk_size <= text_length <= self.strategy.max_chunk_size,
            'has_multiple_types': len(element_types) > 1
        }


# 工具函数
def get_parser_for_file_type(file_type: str):
    """根据文件类型获取解析器类"""
    processors = {
        'pdf': PDFParser if PDF_PARSER_AVAILABLE else None,
        'docx': DocxParser if DOCX_PARSER_AVAILABLE else None,
        'xlsx': ExcelParser,
        'csv': ExcelParser,
        'xmind': XMindParser if XMIND_PARSER_AVAILABLE else None,
        'txt': TextParser,
        'md': TextParser,
        'jpg': ImageParser if OCR_PARSER_AVAILABLE else None,
        'png': ImageParser if OCR_PARSER_AVAILABLE else None,
        'feishu': FeishuParser if FEISHU_PARSER_AVAILABLE else None
    }
    return processors.get(file_type.lower())


# 示例使用
if __name__ == "__main__":
    print("智能切片处理器测试")
    print("=" * 60)
    
    # 创建处理器
    processor = ChunkProcessor(strategy='hybrid')
    
    # 测试数据：模拟元素
    test_elements = [
        ContentElement('heading', '第一章 项目概述', {'position': 0}),
        ContentElement('text', '本项目旨在开发一个智能风险分析系统...' * 10, {'position': 1}),
        ContentElement('text', '系统将包含以下核心功能...' * 8, {'position': 2}),
        ContentElement('heading', '第二章 技术架构', {'position': 3}),
        ContentElement('text', '系统采用微服务架构...' * 12, {'position': 4}),
        ContentElement('table', [['组件', '技术'], ['后端', 'Python'], ['前端', 'React']], {'position': 5}),
        ContentElement('text', '数据库设计如下...' * 6, {'position': 6})
    ]
    
    print(f"测试元素数量: {len(test_elements)}")
    
    # 测试切片
    chunks = processor.strategy.chunk_elements(test_elements)
    
    print(f"生成的切片数量: {len(chunks)}")
    
    for i, chunk in enumerate(chunks):
        chunk_text = processor.chunk_to_text(chunk)[:100] + "..."
        quality = processor.analyze_chunk_quality(chunk)
        
        print(f"\n切片 {i+1}:")
        print(f"  元素数量: {quality['element_count']}")
        print(f"  文本长度: {quality['text_length']}")
        print(f"  元素类型: {quality['element_types']}")
        print(f"  是否优化: {'✅' if quality['is_optimal'] else '⚠️'}")
        print(f"  预览: {chunk_text}")
    
    print("\n✅ 切片处理器测试完成")